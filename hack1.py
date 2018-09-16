import numpy
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.stem import WordNetLemmatizer
import PyPDF2
from gensim.summarization import summarizer
import re
# import ocrmypdf
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
# import pdftotext
import textract
import base64
import docusign_esign as docusign
from docusign_esign import AuthenticationApi, EnvelopesApi, TemplatesApi, DiagnosticsApi
from docusign_esign.rest import ApiException

# import sys
# import os
# import comtypes.client



def extract_text(pdf):
    document = open(pdf, 'rb')
    reader = PyPDF2.PdfFileReader(document)
    # reader = pdftotext.PDF(document)
    numPages = reader.getNumPages()
    allText = ""

    for i in range(numPages):
        page = reader.getPage(i)
        text = page.extractText()
        allText += text

    # if allText == "":
    #   my_ocr = pocr()
    #   allText = my_ocr.go(pdf)
    print(allText)
    return allText

def get_coodinates(text, important):
    # print(text)
    text = text.lower()
    sentences = summarizer.summarize(text, split = True)
    # print(sentences)
    # lemmatizer = WordNetLemmatizer()
    # lematizer.lemmatize(text)
    for s in sent_tokenize(text):
        words = set(word_tokenize(s))
        # w2 = []
        # for w in words:
        #   w2.append(lemmatizer.lemmatize(w))
        # words = set(w2)
        if words.intersection(important):
            print(s)
            print('INTERSECTION')
            sentences.append(s)
    coordinates = []
    for s in sentences:
        coordinates.append((text.find(s), text.find(s)+len(s)))

    return coordinates

# def findAll(st, ch):
#   first = st.find(ch)
#   indices = [first]
#   i = 0
#   while first != -1 and i < len(st[i:]):
#       i = first+1
#       first = st[i:].find(ch)
#       indices.append(first)
#       print("first:", first)
#       print("i:", i)
#   return indices
def findAll(s, ch):
    return [i for i, ltr in enumerate(s) if ltr == ch]

def gen_word_doc(text, coords):
    #print(text)
    doc = Document()
    doc.save('contract.docx')
    # doc.add_heading('yoooooooooooooooo')
    paragraphs = findAll(text, '\n')
    for i in range(len(paragraphs)-1):
        lines = []
        low = paragraphs[i]
        for c in coords:
            if low <= c[0] and c[1] <= paragraphs[i+1]:
                lines.append((low+1, c[0], 0))
                lines.append((c[0], c[1], 1))
                low = c[1]
        if low != paragraphs[i+1]:
            lines.append((low, paragraphs[i+1],0))
        para = doc.add_paragraph()
        print(lines)
        for l in lines:
            if l[2] ==0:
                doc.paragraphs[i].add_run(text[l[0]+1: l[1]]) 
                #para.add_run(text[l[0]: l[1]]) 
                # print(text[l[0]: l[1]])
            else:
                doc.paragraphs[i].add_run(text[l[0]: l[1]]).font.highlight_color = WD_COLOR_INDEX.YELLOW
                # para.add_run(text[l[0]: l[1]]).underline = True
                # print(text[l[0]: l[1]])


    doc.save('contract.docx')
    return doc

def testRequestASignature():
    username = "22a3670e-5f13-45fd-b2b8-ecca4a8e2b25"
    password = "applepie"
    integrator_key = "26ec5e49-a531-4457-ab81-163d76f37a17"
    BASE_URL = "https://demo.docusign.net/restapi"
    user_id = "39c6e30b-3b99-486f-a671-96acc240d7ab"
    oauth_base_url = "account-d.docusign.com" # use account.docusign.com for Live/Production
    api_client = docusign.ApiClient(BASE_URL)
    redirect_uri = "https://www.docusign.com"
    private_key_filename = 'private_key2.txt'
    
    # IMPORTANT NOTE:
    # the first time you ask for a JWT access token, you should grant access by making the following call
    # get DocuSign OAuth authorization url:
    oauth_login_url = api_client.get_jwt_uri(integrator_key, redirect_uri, oauth_base_url)
    # open DocuSign OAuth authorization url in the browser, login and grant access
    # webbrowser.open_new_tab(oauth_login_url)
    print(oauth_login_url)
    # END OF NOTE

    # configure the ApiClient to asynchronously get an access to token and store it
    api_client.configure_jwt_authorization_flow(private_key_filename, oauth_base_url, integrator_key, user_id, 3600)

    docusign.configuration.api_client = api_client

    # sign_test1_file = "test/docs/SignTest1.pdf"
    file_contents = open('/Users/lisayoo/Desktop/contract.pdf', 'rb').read()
  
    # create an envelope to be signed
    envelope_definition = docusign.EnvelopeDefinition()
    envelope_definition.email_subject = "Hi! Here's your annotated contract."
    envelope_definition.email_blurb = 'Hello, Please sign my Python SDK Envelope.'

    # add a document to the envelope
    doc = docusign.Document()
    base64_doc = base64.b64encode(file_contents).decode("utf-8")
    doc.document_base64 = base64_doc
    doc.name = "test.pdf"
    doc.document_id = '1'
    envelope_definition.documents = [doc]

    # Add a recipient to sign the document
    signer = docusign.Signer()
    signer.email = "kwicks@mit.edu"
    signer.name = 'Kat Wicks'
    signer.recipient_id = '1'

    # Create a SignHere tab somewhere on the document for the signer to sign
    # sign_here = docusign.SignHere()
    # sign_here.document_id = '1'
    # sign_here.page_number = '1'
    # sign_here.recipient_id = '1'
    # sign_here.x_position = '100'
    # sign_here.y_position =  '100'
    # sign_here.scale_value = '0.5'

    # tabs = docusign.Tabs()
    # tabs.sign_here_tabs = [sign_here]
    # signer.tabs = tabs

    recipients = docusign.Recipients()
    recipients.signers = [signer]
    envelope_definition.recipients = recipients

    envelope_definition.status = 'sent'

    auth_api = AuthenticationApi()
    envelopes_api = EnvelopesApi()

    try:
        login_info = auth_api.login(api_password='true', include_account_id_guid='true')
        assert login_info is not None
        assert len(login_info.login_accounts) > 0
        login_accounts = login_info.login_accounts
        assert login_accounts[0].account_id is not None

        base_url, _ = login_accounts[0].base_url.split('/v2')
        api_client.host = base_url
        docusign.configuration.api_client = api_client

        envelope_summary = envelopes_api.create_envelope(login_accounts[0].account_id, envelope_definition=envelope_definition)
        assert envelope_summary is not None
        assert envelope_summary.envelope_id is not None
        assert envelope_summary.status == 'sent'

        print("EnvelopeSummary: ", end="")
        print(envelope_summary)

    except ApiException as e:
        print("\nException when calling DocuSign API: %s" % e)
        assert e is None # make the test case fail in case of an API exception

#important_phrases = set(("arbitration", "sign","initial", "$", "probation", "leave", "confidentiality", "non-compete", "termination", "penalty", "liability", "indemnity", "conflict resolution", "damages", "time is of the essence", "nullification", "entire agreement", "fees", "subcontract", "force maejeure", "governed by the laws of"))
important_phrases = set(("arbitration", "signature","initial", "governed", "$", "probation", "leave", "confidentiality", "non-compete", "termination", "penalty", "liability", "indemnity", "conflict", "damages", "nullification",  "fees", "subcontract"))
lemmatizer = WordNetLemmatizer()
contract_words = []
# for i in important_phrases:
#   contract_words.append(lemmatizer.lemmatize(i))
contract_words = set(contract_words)
# doc = extract_text('/Users/lisayoo/Downloads/ccss.pdf')
doc = str(textract.process("/Users/lisayoo/Desktop/business_sample.pdf", layout=True, encoding='ASCII'))
doc = doc.replace("\\n", "\n")
print(str(doc))
# print(doc.count(b'\n'))
# doc = str(doc.decode('ASCII'))
# print(type(doc))
#print(doc)
# print(doc)
to_highlight = get_coodinates(str(doc), important_phrases)
highlighted_doc = gen_word_doc(doc, to_highlight)

# in_file = os.path.abspath('/Desktop/contract.docx')
# out_file = os.path.abspath('/Desktop/contract.pdf')

# word = comtypes.client.CreateObject('/Desktop/contract.pdf')
# doc = word.Documents.Open(in_file)
# doc.SaveAs(out_file, FileFormat=wdFormatPDF)
# doc.Close()
# word.Quit()

testRequestASignature()
#with open("contract.docx", "rb") as f:

    #highlighted_b64 = base64.b64encode(f.read())
#   testRequestASignature(f.read())
